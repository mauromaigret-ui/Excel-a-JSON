from __future__ import annotations

import re
from typing import Dict, List

import pandas as pd
from docx import Document

from app.config import RESULTS_DIR


def _format_n(value: float) -> int | float:
    if value is None:
        return 0
    if abs(value - round(value)) < 1e-6:
        return int(round(value))
    return round(value, 1)


def _safe_filename(name: str) -> str:
    name = name.strip()
    name = name.replace("/", "_")
    name = re.sub(r"[^A-Za-z0-9_ -]+", "", name)
    name = re.sub(r"\s+", "_", name)
    return name[:80] if name else "grupo"


def build_reports(
    var_sum: Dict[str, float],
    group_specs: Dict[str, Dict],
    labels: Dict[str, str],
    output_prefix: str,
) -> Dict[str, object]:
    reports = []

    for group_title, spec in group_specs.items():
        vars_list = spec["variables"]
        if not vars_list:
            continue

        category_col = spec.get("category_col")
        category_map = spec.get("category_map", {})
        label_override = spec.get("label_override", {})
        denominator = spec.get("denominator", "sum")
        total_label = spec.get("total_label", "Total")
        no_total = spec.get("no_total", False)

        rows = []

        for code in vars_list:
            value = float(var_sum.get(code, 0.0))
            value_display = _format_n(value)
            row = {
                "Etiqueta": label_override.get(code, labels.get(code, code)),
                "n": value_display,
                "Porcentaje": "",
                "is_total": False,
                "is_subtotal": False,
            }
            if category_col:
                row[category_col] = category_map.get(code, "")
            rows.append(row)

        if denominator == "by_category" and category_col:
            categories = {}
            for row in rows:
                cat = row.get(category_col, "")
                categories.setdefault(cat, 0.0)
                categories[cat] += float(row["n"]) if row["n"] is not None else 0.0

            for row in rows:
                cat = row.get(category_col, "")
                denom = categories.get(cat, 0.0)
                if denom > 0:
                    row["Porcentaje"] = f"{round((float(row['n']) / denom) * 100, 1)}%"

            for cat, subtotal in categories.items():
                subtotal_display = _format_n(subtotal)
                rows.append(
                    {
                        "Etiqueta": f"{total_label} {cat}",
                        category_col: cat,
                        "n": subtotal_display,
                        "Porcentaje": "100%" if subtotal > 0 else "",
                        "is_total": True,
                        "is_subtotal": True,
                    }
                )
        else:
            if denominator in var_sum:
                denom_value = float(var_sum.get(denominator, 0.0))
            elif denominator == "sum":
                denom_value = sum(float(r["n"]) for r in rows if r["n"] is not None)
            else:
                denom_value = 0.0

            for row in rows:
                if denom_value > 0 and row["n"] is not None:
                    row["Porcentaje"] = f"{round((float(row['n']) / denom_value) * 100, 1)}%"

            if not no_total and total_label:
                total_display = _format_n(denom_value)
                rows.append(
                    {
                        "Etiqueta": total_label,
                        **({category_col: ""} if category_col else {}),
                        "n": total_display,
                        "Porcentaje": "100%" if denom_value > 0 else "",
                        "is_total": True,
                        "is_subtotal": False,
                    }
                )

        # build CSV
        safe_name = _safe_filename(group_title)
        csv_path = RESULTS_DIR / f"{output_prefix}{safe_name}.csv"
        df_rows = pd.DataFrame(rows)
        cols = [c for c in df_rows.columns if c not in {"is_total", "is_subtotal"}]
        if category_col:
            cols = [category_col] + [c for c in cols if c != category_col]
        df_rows[cols].to_csv(csv_path, index=False)

        reports.append(
            {
                "title": group_title,
                "rows": rows,
                "csv_path": str(csv_path),
            }
        )

    # consolidated outputs
    combined_csv = RESULTS_DIR / f"{output_prefix}consolidado.csv"
    combined_xlsx = RESULTS_DIR / f"{output_prefix}consolidado.xlsx"
    combined_html = RESULTS_DIR / f"{output_prefix}consolidado.html"
    combined_docx = RESULTS_DIR / f"{output_prefix}consolidado.docx"

    all_rows = []
    for rep in reports:
        for row in rep["rows"]:
            out = {"Variable": rep["title"], **{k: v for k, v in row.items() if k not in {"is_total", "is_subtotal"}}}
            all_rows.append(out)

    combined_df = pd.DataFrame(all_rows)
    combined_df.to_csv(combined_csv, index=False)

    with pd.ExcelWriter(combined_xlsx, engine="openpyxl") as writer:
        combined_df.to_excel(writer, index=False, sheet_name="Consolidado")
        for rep in reports:
            sheet_name = _safe_filename(rep["title"])[:31]
            df_rows = pd.DataFrame(rep["rows"])
            cols = [c for c in df_rows.columns if c not in {"is_total", "is_subtotal"}]
            df_rows[cols].to_excel(writer, index=False, sheet_name=sheet_name)

    html_parts = ["<h1>Reporte consolidado</h1>"]
    for rep in reports:
        html_parts.append(f"<h2>{rep['title']}</h2>")
        df_rows = pd.DataFrame(rep["rows"])
        cols = [c for c in df_rows.columns if c not in {"is_total", "is_subtotal"}]
        html_parts.append(df_rows[cols].to_html(index=False))
    combined_html.write_text("\n".join(html_parts), encoding="utf-8")

    try:
        doc = Document()
        doc.add_heading("Reporte consolidado", level=1)
        for rep in reports:
            doc.add_heading(rep["title"], level=2)
            df_rows = pd.DataFrame(rep["rows"])
            cols = [c for c in df_rows.columns if c not in {"is_total", "is_subtotal"}]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for idx, col in enumerate(cols):
                run = hdr[idx].paragraphs[0].add_run(col)
                run.bold = True
            for _, row in df_rows.iterrows():
                cells = table.add_row().cells
                is_total = bool(row.get("is_total", False))
                for idx, col in enumerate(cols):
                    text = "" if pd.isna(row[col]) else str(row[col])
                    run = cells[idx].paragraphs[0].add_run(text)
                    if is_total:
                        run.bold = True
        doc.save(combined_docx)
    except Exception:
        combined_docx = RESULTS_DIR / f"{output_prefix}consolidado_docx_error.txt"
        combined_docx.write_text("Error generando DOCX. Use el HTML o XLSX.")

    return {
        "reports": reports,
        "combined_csv": str(combined_csv),
        "combined_html": str(combined_html),
        "combined_docx": str(combined_docx),
        "combined_xlsx": str(combined_xlsx),
    }
